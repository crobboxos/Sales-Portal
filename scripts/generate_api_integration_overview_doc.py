from __future__ import annotations

import csv
import json
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CONFIG_DIR = ROOT / "config"
OUTPUT_DOC = CONFIG_DIR / "Salesforce_API_Integration_Scoping_Overview.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def maybe_add_image(document: Document, image_path: Path, width: Inches, caption: str) -> None:
    if not image_path.exists():
        return
    document.add_picture(str(image_path), width=width)
    add_caption(document, caption)


def short_fields(value: str, max_chars: int = 95) -> str:
    if len(value) <= max_chars:
        return value
    return value[: max_chars - 3] + "..."


def main() -> None:
    org_info = json.loads((CONFIG_DIR / "org-info.json").read_text(encoding="utf-8-sig"))
    objects = read_csv(CONFIG_DIR / "objects.csv")
    relationships = read_csv(CONFIG_DIR / "object-relationships.csv")
    focused_edges = read_csv(CONFIG_DIR / "object-link-summary-focused.csv")
    oppty_edges = read_csv(CONFIG_DIR / "object-link-summary-oppty-flow.csv")

    total_objects = len(objects)
    standard_objects = sum(1 for r in objects if r["category"] == "standard")
    custom_objects = sum(1 for r in objects if r["category"] == "custom")
    custom_mdt = sum(1 for r in objects if r["api_name"].endswith("__mdt"))
    custom_events = sum(1 for r in objects if r["api_name"].endswith("__e"))

    rel_count = len(relationships)
    source_object_count = len({r["source_object"] for r in relationships})
    target_object_count = len({r["target_object"] for r in relationships})

    rel_type_counts = Counter(r["relationship_type"] for r in relationships)
    target_resolution_counts = Counter(r["target_resolution"] for r in relationships)

    custom_prefix_counter: Counter[str] = Counter()
    for r in objects:
        name = r["api_name"]
        if r["category"] != "custom":
            continue
        if "__" in name:
            custom_prefix_counter[name.split("__", 1)[0] + "__"] += 1
        else:
            custom_prefix_counter["Unmanaged"] += 1

    top_prefix_rows = [
        [prefix, str(count)] for prefix, count in custom_prefix_counter.most_common(12)
    ]

    top_focus_sources = Counter(r["source_object"] for r in focused_edges).most_common(8)
    top_focus_targets = Counter(r["target_object"] for r in focused_edges).most_common(8)

    chain_priority = {
        "Opportunity",
        "SCMC__Customer_Quotation__c",
        "SCMC__Sales_Order__c",
        "SCMC__Invoicing__c",
        "c2g__codaInvoice__c",
        "Commission_Payment__c",
        "Account",
        "Contact",
    }
    selected_chain_edges: list[dict[str, str]] = []
    for row in oppty_edges:
        if row["source_object"] in chain_priority or row["target_object"] in chain_priority:
            selected_chain_edges.append(row)
    selected_chain_edges = selected_chain_edges[:18]

    finance_priority = {
        "SCMC__Invoicing__c",
        "c2g__codaInvoice__c",
        "c2g__codaCreditNote__c",
        "c2g__codaTransaction__c",
        "c2g__codaCompany__c",
        "c2g__codaIntercompanyTransfer__c",
        "c2g__codaGeneralLedgerAccount__c",
        "SCMC__Purchase_Order__c",
        "SCMC__Purchase_Order_Line_Item__c",
    }
    selected_finance_edges: list[dict[str, str]] = []
    for row in focused_edges:
        if row["source_object"] in finance_priority or row["target_object"] in finance_priority:
            selected_finance_edges.append(row)
    selected_finance_edges = selected_finance_edges[:18]

    doc = Document()

    # Global style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title page
    title = doc.add_heading("Salesforce API Integration Scoping Overview", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Sales-Portal sandbox baseline and integration blueprint for external system connectivity"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_utc = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    metadata_rows = [
        ["Document purpose", "Scope Salesforce integration approach via APIs for downstream systems"],
        ["Source org alias", org_info["alias"]],
        ["Source org id", org_info["orgId"]],
        ["Source instance URL", org_info["instanceUrl"]],
        ["Salesforce API version", str(org_info["apiVersion"])],
        ["Metadata snapshot timestamp", org_info["retrievedAtUtc"]],
        ["Document generated", generated_utc],
    ]
    add_table(doc, ["Field", "Value"], metadata_rows)

    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document provides a technical scoping baseline for integrating Salesforce with external systems "
        "through API-based patterns. It combines org metadata inventory, object relationship mapping, and "
        "focused business-flow diagrams generated from the current sandbox snapshot."
    )
    for item in [
        f"Org baseline includes {total_objects} objects ({standard_objects} standard, {custom_objects} custom).",
        f"Relationship map includes {rel_count} lookup-style links across {source_object_count} source objects.",
        "Commercial object chain and finance integration paths are mapped and included as ERD diagrams.",
        "Recommended primary integration posture is API-first with OAuth 2.0 server-to-server authentication.",
        "A phased implementation plan is included for delivery, governance, and operational readiness.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Environment and Metadata Baseline", level=1)
    doc.add_paragraph(
        "The following baseline is derived from files generated under config/ in this repository."
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total objects discovered", str(total_objects)],
            ["Standard objects", str(standard_objects)],
            ["Custom objects", str(custom_objects)],
            ["Custom metadata types (__mdt)", str(custom_mdt)],
            ["Custom platform events (__e)", str(custom_events)],
            ["Relationship rows discovered", str(rel_count)],
            ["Objects with outbound relationships", str(source_object_count)],
            ["Unique relationship targets", str(target_object_count)],
            ["Lookup relationships", str(rel_type_counts.get("Lookup", 0))],
            ["RecordType relationships", str(rel_type_counts.get("RecordType", 0))],
            [
                "Target resolution from ReferenceTo",
                str(target_resolution_counts.get("ReferenceTo", 0)),
            ],
            ["Target resolution from DataType parsing", str(target_resolution_counts.get("DataType", 0))],
        ],
    )

    doc.add_paragraph("Top custom package/object namespaces in this org:")
    add_table(doc, ["Namespace prefix", "Object count"], top_prefix_rows)

    doc.add_heading("3. Integration-Relevant Data Model Highlights", level=1)
    doc.add_paragraph(
        "This org has a dense managed-package footprint. Integration design should treat Salesforce as a "
        "hub with multiple domain-specific bounded contexts, especially SCMC and c2g finance objects."
    )
    doc.add_paragraph("Highest edge density in focused map (source objects):")
    for name, count in top_focus_sources:
        doc.add_paragraph(f"{name}: {count} focused target links", style="List Bullet")
    doc.add_paragraph("Highest edge density in focused map (target objects):")
    for name, count in top_focus_targets:
        doc.add_paragraph(f"{name}: {count} inbound focused links", style="List Bullet")
    maybe_add_image(
        doc,
        CONFIG_DIR / "erd-focused-commercial.png",
        Inches(6.2),
        "Figure 0 - Focused ERD: Commercial process landscape",
    )

    doc.add_heading("4. Commercial Flow (Opportunity to Invoice)", level=1)
    doc.add_paragraph(
        "The commercial flow connects CRM and order operations from opportunity through quotation, sales order, "
        "invoicing, and invoice-related finance references."
    )
    add_table(
        doc,
        ["Source object", "Target object", "Link count", "Primary field(s)"],
        [
            [
                r["source_object"],
                r["target_object"],
                r["relationship_field_count"],
                short_fields(r["fields"]),
            ]
            for r in selected_chain_edges
        ],
    )
    maybe_add_image(
        doc,
        CONFIG_DIR / "erd-opportunity-flow.png",
        Inches(6.2),
        "Figure 1 - Focused ERD: Opportunity to Invoicing flow",
    )

    doc.add_heading("5. Finance Integration Lens (SCMC to c2g)", level=1)
    doc.add_paragraph(
        "Finance-related links show where commercial records connect into c2g accounting entities such as "
        "sales invoice, credit note, transaction, intercompany transfer, and company/GL structures."
    )
    add_table(
        doc,
        ["Source object", "Target object", "Link count", "Primary field(s)"],
        [
            [
                r["source_object"],
                r["target_object"],
                r["relationship_field_count"],
                short_fields(r["fields"]),
            ]
            for r in selected_finance_edges
        ],
    )
    maybe_add_image(
        doc,
        CONFIG_DIR / "erd-focused-finance.png",
        Inches(6.2),
        "Figure 2 - Focused ERD: Finance integration flow",
    )

    doc.add_heading("6. API Access Model and Recommended Patterns", level=1)
    doc.add_paragraph(
        "Recommended default is API-first orchestration via a middleware or integration service account. "
        "Use object and relationship metadata to drive payload mapping and contract ownership."
    )
    add_table(
        doc,
        ["Salesforce API", "Best use", "Typical volume/latency", "Implementation notes"],
        [
            [
                "REST API",
                "CRUD, query, orchestration",
                "Low/medium volume, near real-time",
                "Default for synchronous integration flows.",
            ],
            [
                "Composite REST",
                "Bundled writes/reads in one call",
                "Low/medium volume, near real-time",
                "Reduces round-trips; useful for transactional orchestration.",
            ],
            [
                "Bulk API 2.0",
                "Large inserts/updates/extracts",
                "High volume, async",
                "Use for backfill, reconciliation, and nightly integration windows.",
            ],
            [
                "Streaming (Platform Events/CDC)",
                "Event-driven downstream updates",
                "Near real-time",
                "Use where low-latency propagation is required.",
            ],
            [
                "Metadata + Tooling APIs",
                "Schema and deployment automation",
                "N/A",
                "Use in CI/CD and schema governance, not business-transaction payloads.",
            ],
        ],
    )

    doc.add_paragraph("Recommended authentication profile:")
    for item in [
        "Use a dedicated integration user per system or integration domain.",
        "Use Connected App with OAuth 2.0 JWT Bearer flow for server-to-server calls.",
        "Restrict scopes to least privilege and enforce IP/session policies.",
        "Store private keys and client credentials in a secure secret manager.",
        "Separate sandbox and production connected apps/credentials.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Suggested endpoint baseline:")
    doc.add_paragraph(
        f"{org_info['instanceUrl']}/services/data/v{org_info['apiVersion']}/",
        style="Intense Quote",
    )

    doc.add_heading("7. Integration Scoping Matrix", level=1)
    add_table(
        doc,
        ["Use case type", "Direction", "Pattern", "Key Salesforce objects", "Design notes"],
        [
            [
                "Order sync to ERP",
                "Salesforce -> ERP",
                "Event + REST/Composite",
                "SCMC__Sales_Order__c, SCMC__Sales_Order_Line_Item__c, SCMC__Invoicing__c",
                "Publish order state changes and perform idempotent upserts.",
            ],
            [
                "Invoice/credit feedback",
                "ERP -> Salesforce",
                "REST + Bulk for backlog",
                "c2g__codaInvoice__c, c2g__codaCreditNote__c, Commission_Payment__c",
                "Use external IDs and replay-safe processing.",
            ],
            [
                "Master data sync",
                "Bidirectional",
                "REST for delta + Bulk for periodic reconciliation",
                "Account, Contact, SCMC__Item__c, c2g__codaCompany__c",
                "Define source-of-truth per field domain.",
            ],
            [
                "Analytics extraction",
                "Salesforce -> data platform",
                "Bulk API 2.0",
                "Opportunity chain + finance objects",
                "Partition exports and enforce incremental strategy.",
            ],
        ],
    )

    doc.add_heading("8. Security, Governance, and Operational Controls", level=1)
    for item in [
        "Establish data classification for all synced entities and fields.",
        "Implement field-level allowlists, not broad object-level extracts.",
        "Apply integration user profile/permission set with minimum required access.",
        "Use idempotency keys and external IDs to prevent duplicate writes.",
        "Set API monitoring with alerts for failures, retries, and limit consumption.",
        "Define retention and audit strategy for request/response payload logs.",
        "Document runbooks for incident triage, replay, and rollback.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Delivery Plan and Effort Bands", level=1)
    add_table(
        doc,
        ["Phase", "Duration (estimate)", "Primary outputs"],
        [
            [
                "Phase 1 - Discovery and contracts",
                "1-2 weeks",
                "Field-level mapping, ownership model, API contracts, non-functional requirements.",
            ],
            [
                "Phase 2 - Foundation build",
                "2-4 weeks",
                "Connected App setup, auth automation, middleware scaffolding, observability baseline.",
            ],
            [
                "Phase 3 - Core flows",
                "3-6 weeks",
                "Opportunity-to-invoice and finance feedback integrations in sandbox/UAT.",
            ],
            [
                "Phase 4 - Hardening and go-live",
                "2-3 weeks",
                "Performance tests, failure handling, runbooks, production cutover and hypercare.",
            ],
        ],
    )

    doc.add_heading("10. Immediate Next Actions", level=1)
    for item in [
        "Confirm target systems and ownership for each integration lane (ERP, data platform, service tooling, etc.).",
        "Approve canonical key strategy (Salesforce IDs vs external IDs) per object family.",
        "Create Connected App(s) and service principals for non-interactive integration access.",
        "Finalize object and field allowlists for Opportunity, SCMC order flow, and c2g finance entities.",
        "Define integration SLAs (latency, RPO/RTO, retry policy, support window).",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("11. Supporting Artifacts in config/", level=1)
    for item in [
        "org-info.json",
        "objects.csv",
        "object-relationships.csv",
        "object-link-summary.csv",
        "object-link-summary-focused.csv",
        "object-link-summary-oppty-flow.csv",
        "erd-focused-commercial.png/.svg/.mmd",
        "erd-focused-finance.png/.svg/.mmd",
        "erd-opportunity-flow.png/.svg/.mmd",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT_DOC)
    print(f"Wrote {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
