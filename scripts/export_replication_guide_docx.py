from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CONFIG_DIR = ROOT / "config"

INPUT_MD = CONFIG_DIR / "PROJECT_CAPABILITY_REPLICATION_GUIDE.md"
OUTPUT_DOCX = CONFIG_DIR / "PROJECT_CAPABILITY_REPLICATION_GUIDE.docx"


def main() -> None:
    lines = INPUT_MD.read_text(encoding="utf-8-sig").splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    title = doc.add_heading("Project Capability Replication Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Salesforce discovery, mapping, runbook, and documentation bootstrap")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    doc.add_paragraph(f"Generated: {generated}", style="Intense Quote")

    for line in lines:
        txt = line.rstrip()
        if not txt:
            continue
        if txt.startswith("# "):
            doc.add_heading(txt[2:].strip(), level=1)
        elif txt.startswith("## "):
            doc.add_heading(txt[3:].strip(), level=2)
        elif txt.startswith("### "):
            doc.add_heading(txt[4:].strip(), level=3)
        elif txt.startswith("```"):
            # Skip markdown fence markers only; content lines will flow as regular text below.
            continue
        elif txt.startswith("- "):
            doc.add_paragraph(txt[2:].strip(), style="List Bullet")
        elif txt[:2].isdigit() and txt[1] == ".":
            # "1. text"
            doc.add_paragraph(txt[3:].strip(), style="List Number")
        elif txt.startswith("1. ") or txt.startswith("2. ") or txt.startswith("3. ") or txt.startswith("4. ") or txt.startswith("5. ") or txt.startswith("6. ") or txt.startswith("7. ") or txt.startswith("8. ") or txt.startswith("9. "):
            doc.add_paragraph(txt[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(txt)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
