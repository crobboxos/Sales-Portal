from __future__ import annotations

import argparse
import csv
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CONFIG_DIR = ROOT / "config"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    parser = argparse.ArgumentParser(description="Export Apex runbook markdown to DOCX.")
    parser.add_argument("--input-md", default=str(CONFIG_DIR / "apex-custom-runbook-all.md"))
    parser.add_argument("--input-csv", default=str(CONFIG_DIR / "apex-custom-runbook-all.csv"))
    parser.add_argument("--output-docx", default=str(CONFIG_DIR / "apex-custom-runbook-all.docx"))
    args = parser.parse_args()

    input_md = Path(args.input_md)
    input_csv = Path(args.input_csv)
    output_docx = Path(args.output_docx)

    md_lines = input_md.read_text(encoding="utf-8-sig").splitlines()
    rows = read_csv(input_csv)

    total = len(rows)
    class_count = sum(1 for r in rows if r.get("artifact_type") == "Class")
    trigger_count = sum(1 for r in rows if r.get("artifact_type") == "Trigger")
    callout_count = sum(1 for r in rows if r.get("has_callout") == "Y")
    rest_count = sum(1 for r in rows if r.get("has_rest_resource") == "Y")
    invocable_count = sum(1 for r in rows if r.get("has_invocable") == "Y")

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    title = doc.add_heading("Apex Runbook (Custom Namespace - All Artifacts)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Operational behavior map for custom classes and triggers")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Generated", generated],
            ["Total artifacts", str(total)],
            ["Classes", str(class_count)],
            ["Triggers", str(trigger_count)],
            ["Callout indicator artifacts", str(callout_count)],
            ["REST classes", str(rest_count)],
            ["Invocable classes", str(invocable_count)],
        ],
    )

    doc.add_page_break()

    for line in md_lines:
        txt = line.rstrip()
        if not txt:
            continue
        if txt.startswith("# "):
            doc.add_heading(txt[2:].strip(), level=1)
        elif txt.startswith("## "):
            # Break per artifact section to keep readability for large runbook.
            if txt[3:].strip()[:2].isdigit():
                doc.add_page_break()
            doc.add_heading(txt[3:].strip(), level=2)
        elif txt.startswith("- "):
            doc.add_paragraph(txt[2:].strip(), style="List Bullet")
        elif txt.startswith("  - "):
            doc.add_paragraph(txt[4:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(txt)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    print(f"Wrote {output_docx}")


if __name__ == "__main__":
    main()
