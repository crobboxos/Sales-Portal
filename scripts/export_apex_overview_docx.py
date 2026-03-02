from __future__ import annotations

import csv
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CONFIG_DIR = ROOT / "config"

STATS_JSON = CONFIG_DIR / "apex-custom-analysis-stats.json"
RUNBOOK_CSV = CONFIG_DIR / "apex-custom-runbook-all.csv"
TOUCHPOINTS_CSV = CONFIG_DIR / "apex-custom-touchpoints.csv"
OUTPUT_DOCX = CONFIG_DIR / "Apex_API_Integration_Scoping_Overview.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def trim(value: str, max_len: int = 140) -> str:
    if len(value) <= max_len:
        return value
    return value[: max_len - 3] + "..."


def main() -> None:
    stats = json.loads(STATS_JSON.read_text(encoding="utf-8-sig"))
    runbook = read_csv(RUNBOOK_CSV)
    touchpoints = read_csv(TOUCHPOINTS_CSV)

    top_rows = runbook[:25]
    ext_touchpoints = [t for t in touchpoints if t.get("external_candidate") == "Y"][:40]

    trigger_rows = [r for r in runbook if r.get("artifact_type") == "Trigger"]
    trigger_obj_counts: dict[str, int] = {}
    for r in trigger_rows:
        obj = (r.get("primary_object") or "").strip() or "(not explicit)"
        trigger_obj_counts[obj] = trigger_obj_counts.get(obj, 0) + 1
    top_trigger_objs = sorted(trigger_obj_counts.items(), key=lambda x: x[1], reverse=True)[:12]

    doc = Document()

    # Base formatting
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    title = doc.add_heading("Apex API Integration Scoping Overview", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Customer-owned Apex inventory, behavior map, and integration touchpoint assessment"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_utc = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Document generated", generated_utc],
            ["Analysis snapshot", stats.get("generated_utc", "")],
            ["Custom classes analyzed", str(stats["scope"]["custom_classes"])],
            ["Custom triggers analyzed", str(stats["scope"]["custom_triggers"])],
            ["Artifacts with callout indicators", str(stats["findings"]["artifacts_with_callout_indicators"])],
            ["REST classes (@RestResource)", str(stats["findings"]["rest_exposed_classes"])],
            ["Flow-invocable classes", str(stats["findings"]["invocable_classes"])],
            ["External-candidate touchpoints", str(stats["findings"]["external_candidate_touchpoint_rows"])],
        ],
    )

    doc.add_paragraph(
        "Scope note: this document focuses on `NamespacePrefix = null` Apex (customer-owned code). "
        "The org also contains extensive managed-package Apex that should be reviewed separately if needed."
    )

    doc.add_heading("1. Executive Summary", level=1)
    for bullet in [
        "The custom Apex layer includes both synchronous trigger-driven logic and invocable/callout utility classes.",
        "Highest-impact components center on Opportunity/OpportunityLineItem synchronization and API callout helpers.",
        "A small number of classes appear to implement direct external integration entry points and outbound calls.",
        "Trigger footprint is concentrated on Opportunity, OpportunityLineItem, Quote, and Account object families.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("2. Top Ranked Apex Artifacts", level=1)
    add_table(
        doc,
        ["Rank", "Type", "Name", "Score", "Primary Object", "Flags", "Summary"],
        [
            [
                r.get("rank", ""),
                r.get("artifact_type", ""),
                r.get("name", ""),
                r.get("integration_score", ""),
                r.get("primary_object", ""),
                "callout="
                + r.get("has_callout", "")
                + ", rest="
                + r.get("has_rest_resource", "")
                + ", invocable="
                + r.get("has_invocable", ""),
                trim(r.get("summary", ""), 120),
            ]
            for r in top_rows
        ],
    )

    doc.add_heading("3. External/API Touchpoint Candidates", level=1)
    if ext_touchpoints:
        add_table(
            doc,
            ["Type", "Artifact", "Touchpoint Type", "Detail"],
            [
                [
                    t.get("artifact_type", ""),
                    t.get("name", ""),
                    t.get("touchpoint_type", ""),
                    trim(t.get("touchpoint_detail", ""), 130),
                ]
                for t in ext_touchpoints
            ],
        )
    else:
        doc.add_paragraph("No external-candidate touchpoints were detected by static pattern analysis.")

    doc.add_heading("4. Trigger Coverage", level=1)
    doc.add_paragraph("Top trigger objects by number of custom triggers:")
    for obj, count in top_trigger_objs:
        doc.add_paragraph(f"{obj}: {count}", style="List Bullet")

    doc.add_heading("5. Suggested Next Analysis Steps", level=1)
    for bullet in [
        "Trace top-ranked callout classes to Named Credentials, External Credentials, and endpoint governance.",
        "Map top triggers to downstream DML side effects and recursion controls in related handlers.",
        "Correlate invocable Apex classes with active Flow definitions to finalize integration sequence diagrams.",
        "Run targeted code review on non-test classes with high score and high DML/SOQL indicators.",
        "Add test coverage and failure-handling verification for API-touchpoint classes before production integration changes.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("6. Supporting Artifacts", level=1)
    for item in [
        "apex-classes-custom.csv",
        "apex-triggers-custom.csv",
        "apex-custom-code-map.csv",
        "apex-custom-touchpoints.csv",
        "apex-custom-runbook-all.csv",
        "apex-custom-runbook-all.md",
        "apex-custom-api-touchpoint-overview.md",
        "apex-custom-analysis-stats.json",
        "apex-metadata-custom/classes/*.cls",
        "apex-metadata-custom/triggers/*.trigger",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
